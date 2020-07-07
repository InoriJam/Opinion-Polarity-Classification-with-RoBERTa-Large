import docx
from tqdm import tqdm

pbar = tqdm(total=150782)
doc = docx.Document()
doc2 = docx.Document()
count = 0
with open('./dataset/train.raw', 'r') as r:
    while True:
        line = r.readline()
        if not line:
            break
        q = line.rstrip('\n')
        if count <= 75392:
            doc.add_paragraph(q)
        else:
            doc2.add_paragraph(q)
        a = r.readline().rstrip('\n')
        if count <= 75392:
            doc.add_paragraph(a)
        else:
            doc2.add_paragraph(a)
        pbar.update(2)
        count += 2
doc.save('./dataset/train_en_part1.docx')
doc2.save('./dataset/train_en_part2.docx')
pbar.close()