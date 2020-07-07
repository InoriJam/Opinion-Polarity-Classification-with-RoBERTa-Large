import json
import docx
import mafan

filename = "train"
length_arr = []
with open(filename + '.json', "r") as fp:
    doc = docx.Document()
    out_arr = []
    line = fp.readline()
    while line:
        content = json.loads(line)

        text_a = content.get("question")
        text_b = content.get("answer")
        label = content.get("yesno_answer")
        text_a = text_a.replace(' ', '')
        text_b = text_b.replace(' ', '')
        text_a = text_a.replace('\t', '')
        text_b = text_b.replace('\t', '')
        text_a = text_a.replace('?', '？')
        text_b = text_b.replace('?', '？')
        if text_a[-1] != "？":
            text_a = text_a + "？"
        text_a = mafan.simplify(text_a)
        text_b = mafan.simplify(text_b)

        if label == "" and filename != "test":
            print("miss label")
            label = "Yes"
        '''if filename != "test":
            out_line = text_a + '\t' + text_b + '\t' + label
        else:
            out_line = text_a + '\t' + text_b'''
        out_line = [text_a, text_b]
        out_arr.append(out_line)
        length_arr.append(len(text_a + text_b))
        line = fp.readline()
    print(filename + " aver length is:%d", sum(length_arr) / len(length_arr))
    print(filename + " max length is:%d", max(length_arr))
    print(filename + " over 512 length is:%d", len([x for x in length_arr if x > 512]))
    print(filename + " over 256 length is:%d", len([x for x in length_arr if x > 256]))
    print(filename + " over 128 length is:%d", len([x for x in length_arr if x > 128]))
    print(filename + " contains %d lines" % len(out_arr))
    print("######################################")
    for ele in out_arr:
        doc.add_paragraph(ele[0])
        doc.add_paragraph(ele[1])

doc.save('./dataset/' + filename + '.docx')